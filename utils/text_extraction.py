# utils/text_extraction.py - FIXED VERSION FOR STREAMLIT FILE UPLOADS
"""
Enhanced text extraction utilities for various document formats with Streamlit support.
"""
import os
import re
import logging
import time
from typing import Tuple, Any, Dict, List, Optional, Union
from io import BytesIO
import hashlib

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Track available libraries
AVAILABLE_LIBS = {
    'pdf': False,
    'docx': False,
    'tesseract': False,
    'pdf2image': False,
    'langdetect': False,
    'magic': False,
    'pil': False  # Ajout pour PIL/Pillow
}

# Try importing optional dependencies with better error handling
try:
    import pypdf
    AVAILABLE_LIBS['pdf'] = True
    logger.info("pypdf library loaded successfully")
except ImportError:
    logger.warning("pypdf not available. PDF text extraction will be limited.")

try:
    import docx
    AVAILABLE_LIBS['docx'] = True
    logger.info("python-docx library loaded successfully")
except ImportError:
    logger.warning("python-docx not available. DOCX text extraction will be limited.")

try:
    import pytesseract
    AVAILABLE_LIBS['tesseract'] = True
    logger.info("pytesseract library loaded successfully")
except ImportError:
    logger.warning("pytesseract not available. OCR will not be available.")

try:
    from pdf2image import convert_from_bytes
    AVAILABLE_LIBS['pdf2image'] = True
    logger.info("pdf2image library loaded successfully")
except ImportError:
    logger.warning("pdf2image not available. OCR for scanned PDFs will not be available.")

try:
    from langdetect import detect as detect_lang, DetectorFactory
    DetectorFactory.seed = 0  # Ensure consistent results
    AVAILABLE_LIBS['langdetect'] = True
    logger.info("langdetect library loaded successfully")
except ImportError:
    logger.warning("langdetect not available. Using fallback language detection.")

# PIL/Pillow imports - CORRECTION PRINCIPALE
try:
    from PIL import Image, ImageEnhance, ImageFilter
    AVAILABLE_LIBS['pil'] = True
    logger.info("PIL/Pillow library loaded successfully")
except ImportError:
    AVAILABLE_LIBS['pil'] = False
    logger.warning("PIL/Pillow not available. Image processing will be limited.")
    
    # Fallback classes pour éviter les erreurs NameError
    class Image:
        @staticmethod
        def fromarray(array):
            return None
    
    class ImageEnhance:
        @staticmethod
        def Contrast(image):
            class MockEnhancer:
                def enhance(self, factor):
                    return image
            return MockEnhancer()
        
        @staticmethod  
        def Sharpness(image):
            class MockEnhancer:
                def enhance(self, factor):
                    return image
            return MockEnhancer()
    
    class ImageFilter:
        @staticmethod
        def MedianFilter():
            return None

# Configuration constants
MAX_OCR_PAGES = 20
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
CHUNK_SIZE = 1024 * 1024  # 1MB chunks for processing
SUPPORTED_ENCODINGS = ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1', 'cp1252']

class DocumentProcessor:
    """Enhanced document processor with Streamlit file upload support"""
    
    def __init__(self, enable_ocr: bool = True, max_pages: int = None):
        self.enable_ocr = enable_ocr
        self.max_pages = max_pages or MAX_OCR_PAGES
        self.processing_stats = {}
        
    def extract_text_pdf_advanced(self, file_content: bytes, use_ocr: bool = True) -> Tuple[str, Dict[str, Any]]:
        """
        Advanced PDF text extraction with metadata and error recovery
        """
        if not AVAILABLE_LIBS['pdf']:
            return "PDF text extraction not available. Install pypdf.", {'error': 'pypdf_missing'}
        
        extracted_text = ""
        metadata = {
            'pages': 0,
            'method': 'direct',
            'confidence': 0.0,
            'processing_time': 0,
            'file_size': len(file_content),
            'errors': []
        }
        
        start_time = time.time()
        
        try:
            # First try direct text extraction
            pdf = pypdf.PdfReader(BytesIO(file_content))
            metadata['pages'] = len(pdf.pages)
            
            # Extract metadata from PDF
            if pdf.metadata:
                try:
                    metadata.update({
                        'title': str(pdf.metadata.get('/Title', '')),
                        'author': str(pdf.metadata.get('/Author', '')),
                        'subject': str(pdf.metadata.get('/Subject', '')),
                        'creator': str(pdf.metadata.get('/Creator', '')),
                        'producer': str(pdf.metadata.get('/Producer', ''))
                    })
                except Exception as e:
                    metadata['errors'].append(f"Metadata extraction error: {str(e)}")
            
            # Extract text from each page
            for page_num in range(min(len(pdf.pages), self.max_pages or len(pdf.pages))):
                try:
                    page = pdf.pages[page_num]
                    page_text = page.extract_text() or ""
                    extracted_text += page_text + "\n\n"
                except Exception as e:
                    metadata['errors'].append(f"Page {page_num + 1}: {str(e)}")
                    continue
            
            # Calculate confidence based on text density
            if extracted_text.strip():
                avg_text_per_page = len(extracted_text) / max(metadata['pages'], 1)
                metadata['confidence'] = min(avg_text_per_page / 500, 1.0)
            
            # Check if we got meaningful text
            if len(extracted_text.strip()) < 50 and use_ocr and self.enable_ocr:
                logger.info("Minimal text extracted from PDF. Trying OCR...")
                ocr_text, ocr_metadata = self.extract_text_pdf_ocr_advanced(file_content)
                if len(ocr_text.strip()) > len(extracted_text.strip()):
                    extracted_text = ocr_text
                    metadata.update(ocr_metadata)
                    metadata['method'] = 'ocr'
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            metadata['errors'].append(f"PDF extraction error: {str(e)}")
            
            # Attempt OCR as fallback
            if use_ocr and self.enable_ocr and AVAILABLE_LIBS['tesseract'] and AVAILABLE_LIBS['pdf2image']:
                logger.info("Trying OCR as fallback...")
                try:
                    ocr_text, ocr_metadata = self.extract_text_pdf_ocr_advanced(file_content)
                    extracted_text = ocr_text
                    metadata.update(ocr_metadata)
                    metadata['method'] = 'ocr_fallback'
                except Exception as ocr_e:
                    extracted_text = f"Error extracting text: {str(e)}. OCR also failed: {str(ocr_e)}"
            else:
                extracted_text = f"Error extracting text: {str(e)}"
        
        metadata['processing_time'] = time.time() - start_time
        return extracted_text, metadata

    def extract_text_pdf_ocr_advanced(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Méthode OCR avancée corrigée"""
        if not AVAILABLE_LIBS['pdf2image'] or not AVAILABLE_LIBS['tesseract']:
            return "OCR not available. Missing pdf2image or pytesseract.", {'error': 'ocr_dependencies_missing'}
        
        extracted_text = ""
        metadata = {
            'pages_processed': 0,
            'ocr_confidence': 0.0,
            'errors': []
        }
        
        try:
            images = convert_from_bytes(file_content, dpi=200)
            
            for i, image in enumerate(images[:self.max_pages]):
                try:
                    # Utiliser OCR de base si PIL n'est pas disponible
                    if AVAILABLE_LIBS['pil']:
                        processed_image = self._preprocess_image_for_ocr(image)
                    else:
                        processed_image = image
                    
                    # OCR avec configuration de base
                    text = pytesseract.image_to_string(processed_image, lang='fra+eng')
                    extracted_text += text + "\n\n"
                    metadata['pages_processed'] += 1
                    
                except Exception as e:
                    metadata['errors'].append(f"Page {i + 1}: {str(e)}")
                    continue
            
            return extracted_text, metadata
            
        except Exception as e:
            return f"OCR failed: {e}", {'error': str(e)}

    def extract_text_pdf_ocr_optimized(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """OCR optimisé avec préprocessing d'image - Version corrigée"""
        if not AVAILABLE_LIBS['pdf2image'] or not AVAILABLE_LIBS['tesseract']:
            return "OCR optimisé non disponible. Dépendances manquantes.", {'error': 'ocr_dependencies_missing'}
        
        try:
            images = convert_from_bytes(file_content, dpi=300)
            extracted_text = ""
            metadata = {'pages_processed': 0, 'preprocessing_applied': []}
            
            for i, image in enumerate(images[:self.max_pages]):
                # Préprocessing seulement si PIL est disponible
                if AVAILABLE_LIBS['pil']:
                    processed_image = self._preprocess_image_for_ocr(image)
                    metadata['preprocessing_applied'].append(f"Page {i+1}: PIL preprocessing")
                else:
                    processed_image = image
                    metadata['preprocessing_applied'].append(f"Page {i+1}: No preprocessing (PIL unavailable)")
                
                # OCR avec configuration optimisée
                custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyzÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞßàáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿ.,;:!?()-[]{}"\' '
                
                text = pytesseract.image_to_string(processed_image, config=custom_config, lang='fra+eng+deu')
                extracted_text += text + "\n\n"
                metadata['pages_processed'] += 1
            
            return extracted_text, metadata
            
        except Exception as e:
            return f"OCR optimisé échoué: {e}", {'error': str(e)}

    def _preprocess_image_for_ocr(self, image) -> object:
        """Préprocessing d'image pour améliorer l'OCR - Version corrigée"""
        # Vérifier si PIL est disponible
        if not AVAILABLE_LIBS['pil']:
            logger.warning("PIL not available, returning image without preprocessing")
            return image
        
        try:
            import cv2
            import numpy as np
            
            # Convertir en numpy array
            img_array = np.array(image)
            
            # Convertir en niveaux de gris
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # Débruitage
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Amélioration du contraste
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(denoised)
            
            # Binarisation adaptative
            binary = cv2.adaptiveThreshold(enhanced, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            # Reconvertir en PIL Image
            return Image.fromarray(binary)
        
        except Exception as e:
            logger.warning(f"Advanced preprocessing failed: {e}, using basic PIL enhancement")
            # Fallback: amélioration basique avec PIL
            try:
                enhanced = ImageEnhance.Contrast(image).enhance(1.5)
                enhanced = ImageEnhance.Sharpness(enhanced).enhance(1.2)
                return enhanced.filter(ImageFilter.MedianFilter())
            except Exception as e2:
                logger.warning(f"Basic PIL enhancement also failed: {e2}, returning original image")
                return image

    def extract_text_docx_advanced(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Advanced DOCX text extraction with structure preservation"""
        if not AVAILABLE_LIBS['docx']:
            return "DOCX text extraction not available. Install python-docx.", {'error': 'docx_missing'}
        
        metadata = {
            'paragraphs': 0,
            'tables': 0,
            'errors': []
        }
        
        start_time = time.time()
        
        try:
            doc = docx.Document(BytesIO(file_content))
            extracted_text = ""
            
            # Extract core properties
            if doc.core_properties:
                try:
                    metadata.update({
                        'title': doc.core_properties.title or '',
                        'author': doc.core_properties.author or '',
                        'subject': doc.core_properties.subject or ''
                    })
                except Exception as e:
                    metadata['errors'].append(f"Properties error: {str(e)}")
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                try:
                    para_text = para.text.strip()
                    if para_text:
                        # Add style markers for headings
                        if 'Heading' in para.style.name:
                            extracted_text += f"\n## {para_text}\n"
                        else:
                            extracted_text += para_text + "\n"
                        metadata['paragraphs'] += 1
                except Exception as e:
                    metadata['errors'].append(f"Paragraph error: {str(e)}")
                    continue
            
            # Extract text from tables
            for table_idx, table in enumerate(doc.tables):
                try:
                    extracted_text += f"\n--- Table {table_idx + 1} ---\n"
                    
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            row_text.append(cell_text)
                        extracted_text += " | ".join(row_text) + "\n"
                    
                    extracted_text += "--- End Table ---\n\n"
                    metadata['tables'] += 1
                    
                except Exception as e:
                    metadata['errors'].append(f"Table {table_idx + 1} error: {str(e)}")
                    continue
                
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            metadata['errors'].append(f"DOCX extraction error: {str(e)}")
            extracted_text = f"Error extracting text: {str(e)}"
        
        metadata['processing_time'] = time.time() - start_time
        metadata['confidence'] = 0.95 if not metadata['errors'] else 0.7
        
        return extracted_text, metadata

    def extract_text_txt_advanced(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Advanced text file extraction with encoding detection"""
        metadata = {
            'encoding': 'unknown',
            'lines': 0,
            'confidence': 0.0,
            'errors': []
        }
        
        start_time = time.time()
        
        try:
            # Try to detect encoding using chardet if available
            detected_encoding = None
            try:
                import chardet
                detection = chardet.detect(file_content)
                if detection['confidence'] > 0.7:
                    detected_encoding = detection['encoding']
                    metadata['encoding'] = detected_encoding
                    metadata['confidence'] = detection['confidence']
            except ImportError:
                pass
            
            # Try different encodings
            encodings_to_try = [detected_encoding] + SUPPORTED_ENCODINGS if detected_encoding else SUPPORTED_ENCODINGS
            
            extracted_text = None
            for encoding in encodings_to_try:
                if encoding is None:
                    continue
                    
                try:
                    extracted_text = file_content.decode(encoding)
                    metadata['encoding'] = encoding
                    if encoding == detected_encoding:
                        metadata['confidence'] = 0.95
                    else:
                        metadata['confidence'] = 0.8
                    break
                except UnicodeDecodeError:
                    continue
            
            # Fallback to latin-1 which should always work
            if extracted_text is None:
                extracted_text = file_content.decode('latin-1', errors='replace')
                metadata['encoding'] = 'latin-1'
                metadata['confidence'] = 0.5
                metadata['errors'].append("Used fallback encoding with error replacement")
            
            # Count lines and basic statistics
            lines = extracted_text.split('\n')
            metadata['lines'] = len(lines)
            metadata['characters'] = len(extracted_text)
            metadata['words'] = len(extracted_text.split())
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            metadata['errors'].append(f"Text extraction error: {str(e)}")
            extracted_text = f"Error extracting text: {str(e)}"
            metadata['confidence'] = 0.0
        
        metadata['processing_time'] = time.time() - start_time
        return extracted_text, metadata

    def detect_language_advanced(self, text: str) -> Tuple[str, Dict[str, Any]]:
        """Advanced language detection with confidence scoring"""
        metadata = {
            'method': 'fallback',
            'confidence': 0.0,
            'candidates': [],
            'text_length': len(text)
        }
        
        if not text or len(text.strip()) < 10:
            return "fr", metadata  # Default to French
        
        # Use langdetect if available
        if AVAILABLE_LIBS['langdetect']:
            try:
                from langdetect import detect_langs
                
                # Detect multiple language candidates
                candidates = detect_langs(text[:5000])  # Use first 5000 chars
                
                if candidates:
                    primary_lang = candidates[0]
                    
                    # Map to supported languages
                    lang_map = {
                        'fr': 'fr', 'en': 'en', 'de': 'de', 'es': 'es',
                        'it': 'fr', 'nl': 'en', 'pt': 'es', 'ca': 'es'
                    }
                    
                    detected_lang = lang_map.get(primary_lang.lang, 'fr')
                    
                    metadata.update({
                        'method': 'langdetect',
                        'confidence': primary_lang.prob,
                        'candidates': [(c.lang, c.prob) for c in candidates[:3]]
                    })
                    
                    return detected_lang, metadata
                    
            except Exception as e:
                logger.warning(f"Language detection failed: {e}")
                metadata['errors'] = [str(e)]
        
        # Fallback to keyword-based detection
        detected_lang, fallback_confidence = self.fallback_language_detection_advanced(text)
        metadata.update({
            'method': 'keyword_frequency',
            'confidence': fallback_confidence
        })
        
        return detected_lang, metadata

    def fallback_language_detection_advanced(self, text: str) -> Tuple[str, float]:
        """Advanced fallback language detection using keyword frequency"""
        text_lower = text.lower()
        
        # Enhanced language markers with weights
        language_markers = {
            "fr": {
                "keywords": {
                    "le": 3, "la": 3, "les": 3, "est": 2, "dans": 2, "pour": 2,
                    "cette": 2, "ils": 2, "nous": 2, "vous": 2, "avec": 2,
                    "sont": 2, "être": 2, "avoir": 2, "qui": 2, "que": 2,
                    "sur": 1, "par": 1, "mais": 1, "ses": 1, "ces": 1
                },
                "patterns": [r'\bdu\b', r'\bdes\b', r'\bau\b', r'\baux\b', r'ç[a-z]', r'[a-z]tion\b']
            },
            "en": {
                "keywords": {
                    "the": 3, "is": 2, "are": 2, "that": 2, "this": 2,
                    "with": 2, "for": 2, "from": 2, "have": 2, "will": 2,
                    "they": 2, "you": 2, "and": 3, "or": 1, "but": 1,
                    "was": 2, "were": 2, "been": 1, "has": 2, "had": 2
                },
                "patterns": [r'\bto\b', r'\bof\b', r'\bin\b', r'\bon\b', r'ing\b', r'ed\b']
            },
            "de": {
                "keywords": {
                    "der": 3, "die": 3, "das": 3, "ist": 2, "und": 3,
                    "für": 2, "ein": 2, "eine": 2, "nicht": 2, "mit": 2,
                    "sich": 1, "auf": 1, "sie": 2, "werden": 2, "wird": 2
                },
                "patterns": [r'\bden\b', r'\bdem\b', r'\bdes\b', r'ung\b', r'keit\b']
            },
            "es": {
                "keywords": {
                    "el": 3, "la": 3, "los": 3, "las": 2, "es": 2,
                    "por": 2, "para": 2, "como": 2, "con": 2, "este": 2,
                    "esta": 2, "que": 3, "se": 2, "del": 2, "en": 2
                },
                "patterns": [r'\bdel\b', r'\bde\b', r'\bal\b', r'ción\b', r'dad\b']
            }
        }
        
        scores = {}
        total_words = len(text_lower.split())
        
        for lang, markers in language_markers.items():
            score = 0
            
            # Count keywords with weights
            for keyword, weight in markers["keywords"].items():
                pattern = r'\b' + re.escape(keyword) + r'\b'
                matches = len(re.findall(pattern, text_lower))
                score += matches * weight
            
            # Count patterns
            for pattern in markers.get("patterns", []):
                matches = len(re.findall(pattern, text_lower))
                score += matches * 0.5
            
            # Normalize score by text length
            if total_words > 0:
                scores[lang] = score / total_words
            else:
                scores[lang] = 0
        
        # Find the best match
        if max(scores.values()) > 0:
            best_lang = max(scores, key=scores.get)
            confidence = min(scores[best_lang] * 10, 1.0)
            return best_lang, confidence
        
        return "fr", 0.1  # Default with low confidence

def extract_text(file, detect_lang: bool = True, **options) -> Tuple[str, str]:
    """
    FIXED main text extraction function with proper Streamlit file upload support
    
    Args:
        file: Streamlit uploaded file or file path
        detect_lang: Whether to detect language from extracted text
        **options: Additional processing options
        
    Returns:
        Tuple of (extracted_text, language_code)
    """
    processor = DocumentProcessor(
        enable_ocr=options.get('enable_ocr', True),
        max_pages=options.get('max_pages')
    )
    
    language = "fr"  # Default language
    extracted_text = ""
    
    try:
        # Handle Streamlit uploaded file vs file path
        if hasattr(file, 'name') and hasattr(file, 'read'):
            # This is a Streamlit uploaded file
            file_name = file.name
            file_content = file.getvalue()  # Use getvalue() for Streamlit files
            logger.info(f"Processing Streamlit uploaded file: {file_name}")
        elif isinstance(file, str):
            # This is a file path
            file_name = os.path.basename(file)
            with open(file, 'rb') as f:
                file_content = f.read()
            logger.info(f"Processing file from path: {file_name}")
        else:
            # Try to handle other file-like objects
            try:
                file_name = getattr(file, 'name', 'unknown')
                if hasattr(file, 'name') and hasattr(file, 'getvalue'):
                    file_content = file.getvalue()
                    file_name = file.name
                elif hasattr(file, 'read'):
                    file_content = file.read()
                    file_name = getattr(file, 'name', 'unknown')
                else:
                    raise ValueError("Unknown file object type")
                logger.info(f"Processing file-like object: {file_name}")
            except Exception as e:
                logger.error(f"Could not read file object: {e}")
                return f"Error reading file: {str(e)}", "fr"
        
        # Check file size
        if len(file_content) > MAX_FILE_SIZE:
            return f"File too large: {len(file_content)} bytes (max: {MAX_FILE_SIZE})", language
        
        # Detect file type from extension
        extension = os.path.splitext(file_name)[1].lower() if file_name else ""
        
        # Extract text based on file type
        if extension == '.pdf':
            extracted_text, extract_metadata = processor.extract_text_pdf_advanced(
                file_content, 
                options.get('use_ocr', True)
            )
        elif extension == '.docx':
            extracted_text, extract_metadata = processor.extract_text_docx_advanced(file_content)
        elif extension == '.txt':
            extracted_text, extract_metadata = processor.extract_text_txt_advanced(file_content)
        else:
            # Try to detect content type from file content if no extension
            if file_content.startswith(b'%PDF'):
                extracted_text, extract_metadata = processor.extract_text_pdf_advanced(
                    file_content, 
                    options.get('use_ocr', True)
                )
            elif file_content.startswith(b'PK'):  # ZIP-based format (DOCX)
                extracted_text, extract_metadata = processor.extract_text_docx_advanced(file_content)
            else:
                # Try as text file
                extracted_text, extract_metadata = processor.extract_text_txt_advanced(file_content)
        
        # Detect language if requested and text is available
        if detect_lang and extracted_text and len(extracted_text.strip()) > 10:
            try:
                language, lang_metadata = processor.detect_language_advanced(extracted_text)
                logger.info(f"Detected language: {language} (confidence: {lang_metadata.get('confidence', 0):.2f})")
            except Exception as e:
                logger.warning(f"Language detection failed: {e}")
                language = "fr"  # Default fallback
        
        # Log processing statistics
        if extract_metadata and 'processing_time' in extract_metadata:
            logger.info(f"Text extraction completed in {extract_metadata['processing_time']:.2f}s")
            logger.info(f"Extracted {len(extracted_text)} characters from {file_name}")
        
        return extracted_text, language
        
    except Exception as e:
        logger.error(f"Error in text extraction: {e}")
        return f"Error extracting text: {str(e)}", "fr"

def validate_file_upload(file) -> Tuple[bool, str]:
    """
    Validate Streamlit file upload
    
    Args:
        file: Streamlit uploaded file
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    if file is None:
        return False, "No file uploaded"
    
    try:
        # Check file size
        file_size = len(file.getvalue())
        if file_size > MAX_FILE_SIZE:
            return False, f"File too large: {file_size / (1024*1024):.1f}MB (max: {MAX_FILE_SIZE/(1024*1024):.0f}MB)"
        
        # Check file extension
        if hasattr(file, 'name'):
            extension = os.path.splitext(file.name)[1].lower()
            supported_extensions = ['.pdf', '.docx', '.txt']
            if extension not in supported_extensions:
                return False, f"Unsupported file type: {extension}. Supported: {', '.join(supported_extensions)}"
        
        # Basic content validation
        content = file.getvalue()
        if len(content) == 0:
            return False, "File is empty"
        
        return True, "File is valid"
        
    except Exception as e:
        return False, f"Error validating file: {str(e)}"

def get_text_statistics(text: str) -> Dict[str, Any]:
    """Get comprehensive text statistics"""
    if not text:
        return {}
    
    lines = text.split('\n')
    words = text.split()
    sentences = re.split(r'[.!?]+', text)
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    return {
        'characters': len(text),
        'characters_no_spaces': len(text.replace(' ', '')),
        'words': len(words),
        'lines': len(lines),
        'sentences': len([s for s in sentences if s.strip()]),
        'paragraphs': len(paragraphs),
        'average_word_length': sum(len(word) for word in words) / len(words) if words else 0,
        'average_sentence_length': len(words) / len([s for s in sentences if s.strip()]) if sentences else 0,
        'reading_time_minutes': len(words) / 200 if words else 0  # Average reading speed
    }

def count_tokens(text: str) -> int:
    """Enhanced token counting with better estimation"""
    if not text:
        return 0
    
    words = text.split()
    punctuation_tokens = len(re.findall(r'[^\w\s]', text)) * 0.5
    return len(words) + int(punctuation_tokens)

if __name__ == "__main__":
    # Test the extraction functions
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        
        print("🔍 Testing enhanced text extraction...")
        print("=" * 50)
        
        text, lang = extract_text(file_path, detect_lang=True)
        
        print(f"✅ Extraction completed")
        print(f"🌍 Detected language: {lang}")
        print(f"📄 Text length: {len(text)} characters")
        
        stats = get_text_statistics(text)
        print(f"📊 Text statistics: {stats}")
        
        print(f"\n📖 Text excerpt:\n{text[:500]}...")
    else:
        print("Please provide a file path as an argument")
        print("Example: python text_extraction.py document.pdf")